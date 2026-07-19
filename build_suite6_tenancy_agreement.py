from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

WORKSPACE = Path("/Users/philipboakye/Documents/OneRoot Expense Register")
OUTPUT_PATH = WORKSPACE / "outputs" / "Suite_6_Patience_Tenancy_Agreement_Once_Paid.docx"
SKILL_ROOT = Path(
    "/Users/philipboakye/.codex/plugins/cache/openai-primary-runtime/documents/26.623.12021/skills/documents"
)
TABLE_GEOMETRY_PATH = SKILL_ROOT / "scripts" / "table_geometry.py"

LANDLORD_NAME = "Philip Boakye"
LANDLORD_ADDRESS_LINES = [
    "P.O. Box CT 10792, Cantonments, Accra",
    "Medie New City",
    "Digital Address: GW-0332-7210",
    "Mobile: 0244620860",
]
PROPERTY_NAME = "Suite 6 - Patience"
PROPERTY_LOCATION = "Medie New City (Parks and Gardens), Accra, Ghana"
RENT_PER_MONTH = "GHS 800"
ADVANCE_RENT = "GHS 9,600"
SERVICE_FEE_TOTAL = "GHS 200"
MOMO_NUMBER = "0242847065"

COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
COLOR_BLUE = RGBColor(0x2E, 0x74, 0xB5)
COLOR_DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
COLOR_MUTED = RGBColor(0x52, 0x60, 0x6D)
COLOR_HEADER_FILL = "F2F4F7"
CONTENT_WIDTH_DXA = 9360
STANDARD_LABEL_DETAIL_WIDTHS_DXA = [2700, 6660]
TWO_UP_WIDTHS_DXA = [4680, 4680]


def load_table_geometry_module():
    spec = importlib.util.spec_from_file_location("table_geometry", TABLE_GEOMETRY_PATH)
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    spec.loader.exec_module(module)
    return module


table_geometry = load_table_geometry_module()


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")


def set_cell_borders(cell, color: str = "D6DDD5", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):
        element = tc_borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def apply_run_font(run, *, name: str = "Calibri", size: int = 11, color=COLOR_BLACK, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def set_paragraph_spacing(paragraph, *, before=0, after=0, line_spacing=1.1, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line_spacing
    paragraph.alignment = alignment


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.start_type = WD_SECTION_START.NEW_PAGE

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLOR_BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.1

    heading_1 = doc.styles["Heading 1"]
    heading_1.font.name = "Calibri"
    heading_1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading_1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading_1.font.size = Pt(16)
    heading_1.font.color.rgb = COLOR_BLUE
    heading_1.font.bold = True
    heading_1.paragraph_format.space_before = Pt(16)
    heading_1.paragraph_format.space_after = Pt(8)
    heading_1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    heading_1.paragraph_format.line_spacing = 1.1

    heading_2 = doc.styles["Heading 2"]
    heading_2.font.name = "Calibri"
    heading_2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading_2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading_2.font.size = Pt(13)
    heading_2.font.color.rgb = COLOR_BLUE
    heading_2.font.bold = True
    heading_2.paragraph_format.space_before = Pt(12)
    heading_2.paragraph_format.space_after = Pt(6)
    heading_2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    heading_2.paragraph_format.line_spacing = 1.1

    heading_3 = doc.styles["Heading 3"]
    heading_3.font.name = "Calibri"
    heading_3._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading_3._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading_3.font.size = Pt(12)
    heading_3.font.color.rgb = COLOR_DARK_BLUE
    heading_3.font.bold = True
    heading_3.paragraph_format.space_before = Pt(8)
    heading_3.paragraph_format.space_after = Pt(4)
    heading_3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    heading_3.paragraph_format.line_spacing = 1.1


def add_centered_line(doc: Document, text: str, *, size: int, color, bold=False, italic=False, after=0):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=0, after=after, line_spacing=1.1, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    run = paragraph.add_run(text)
    apply_run_font(run, size=size, color=color, bold=bold, italic=italic)
    return paragraph


def add_body_paragraph(doc: Document, text: str, *, before=0, after=6, italic=False):
    paragraph = doc.add_paragraph(style="Normal")
    set_paragraph_spacing(paragraph, before=before, after=after, line_spacing=1.1)
    run = paragraph.add_run(text)
    apply_run_font(run, size=11, color=COLOR_BLACK, italic=italic)
    return paragraph


def add_label_value_paragraph(doc: Document, label: str, value: str):
    paragraph = doc.add_paragraph(style="Normal")
    set_paragraph_spacing(paragraph, before=0, after=4, line_spacing=1.1)
    label_run = paragraph.add_run(f"{label}: ")
    apply_run_font(label_run, size=11, color=COLOR_BLACK, bold=True)
    value_run = paragraph.add_run(value)
    apply_run_font(value_run, size=11, color=COLOR_BLACK)
    return paragraph


def add_snapshot_table(doc: Document):
    rows = [
        ("Property", PROPERTY_NAME),
        ("Location", PROPERTY_LOCATION),
        ("Lease Term", "One (1) year from the agreed commencement date to the matching expiry date."),
        ("Rent", f"{RENT_PER_MONTH} per month, payable one year in advance."),
        ("Advance Rent Due Before Activation", ADVANCE_RENT),
        ("Monthly Service Charges", f"{SERVICE_FEE_TOTAL} total per month, split as shown in clause 4."),
        ("Monthly Charges Payment Channel", f"MTN MoMo to {MOMO_NUMBER}"),
        ("Agreement Status", "Prepared for signature once the full advance rent has been paid and confirmed."),
    ]

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table_geometry.apply_table_geometry(
        table,
        STANDARD_LABEL_DETAIL_WIDTHS_DXA,
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=120,
    )

    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(row.cells[0], COLOR_HEADER_FILL)
        for cell in row.cells:
            set_cell_borders(cell)

        p_label = row.cells[0].paragraphs[0]
        p_label.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p_label, before=0, after=0, line_spacing=1.1)
        p_label_run = p_label.add_run(label)
        apply_run_font(p_label_run, size=11, color=COLOR_BLACK, bold=True)

        p_value = row.cells[1].paragraphs[0]
        p_value.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p_value, before=0, after=0, line_spacing=1.1)
        p_value_run = p_value.add_run(value)
        apply_run_font(p_value_run, size=11, color=COLOR_BLACK)


def add_fees_table(doc: Document):
    rows = [
        ("Water and toilet fees", "GHS 60"),
        ("Sweeping of compound and gutter cleaning", "GHS 125"),
        ("Waste disposal", "GHS 15"),
        ("Total monthly service charges", SERVICE_FEE_TOTAL),
    ]

    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"
    table_geometry.apply_table_geometry(
        table,
        [6660, 2700],
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=120,
    )

    headers = ("Charge", "Amount")
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, COLOR_HEADER_FILL)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if header == "Charge" else WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, before=0, after=0, line_spacing=1.1)
        run = p.add_run(header)
        apply_run_font(run, size=11, color=COLOR_BLACK, bold=True)

    for row, (charge, amount) in zip(table.rows[1:], rows):
        for cell in row.cells:
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        p_charge = row.cells[0].paragraphs[0]
        set_paragraph_spacing(p_charge, before=0, after=0, line_spacing=1.1)
        p_charge_run = p_charge.add_run(charge)
        apply_run_font(p_charge_run, size=11, color=COLOR_BLACK, bold=charge.startswith("Total"))

        p_amount = row.cells[1].paragraphs[0]
        p_amount.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p_amount, before=0, after=0, line_spacing=1.1)
        p_amount_run = p_amount.add_run(amount)
        apply_run_font(p_amount_run, size=11, color=COLOR_BLACK, bold=charge.startswith("Total"))


def add_payment_confirmation_table(doc: Document):
    rows = [
        ("Amount received", ADVANCE_RENT),
        ("Payment date", "________________________________________"),
        ("Payment method", "________________________________________"),
        ("Reference / transaction ID", "________________________________________"),
        ("Received by", "________________________________________"),
        ("Agreement effective from", "________________________________________"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table_geometry.apply_table_geometry(
        table,
        STANDARD_LABEL_DETAIL_WIDTHS_DXA,
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=120,
    )
    for row, (label, value) in zip(table.rows, rows):
        for cell in row.cells:
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(row.cells[0], COLOR_HEADER_FILL)

        p_label = row.cells[0].paragraphs[0]
        set_paragraph_spacing(p_label, before=0, after=0, line_spacing=1.1)
        label_run = p_label.add_run(label)
        apply_run_font(label_run, size=11, color=COLOR_BLACK, bold=True)

        p_value = row.cells[1].paragraphs[0]
        set_paragraph_spacing(p_value, before=0, after=0, line_spacing=1.1)
        value_run = p_value.add_run(value)
        apply_run_font(value_run, size=11, color=COLOR_BLACK)


def add_signature_table(doc: Document):
    table = doc.add_table(rows=4, cols=2)
    table_geometry.apply_table_geometry(
        table,
        TWO_UP_WIDTHS_DXA,
        table_width_dxa=CONTENT_WIDTH_DXA,
        indent_dxa=120,
    )

    signatures = [
        ("Landlord Signature", "Tenant Signature"),
        ("_______________________________", "_______________________________"),
        (f"{LANDLORD_NAME}", "[Tenant full name]"),
        ("Date: _________________________", "Date: _________________________"),
    ]

    for row, values in zip(table.rows, signatures):
        for cell, value in zip(row.cells, values):
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, before=0, after=0, line_spacing=1.1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            apply_run_font(run, size=11, color=COLOR_BLACK, bold=row is table.rows[0])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.find(qn("w:tcBorders"))
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            for edge in ("top", "left", "bottom", "right"):
                element = tc_borders.find(qn(f"w:{edge}"))
                if element is None:
                    element = OxmlElement(f"w:{edge}")
                    tc_borders.append(element)
                element.set(qn("w:val"), "nil")


def build_document() -> None:
    doc = Document()
    configure_styles(doc)

    add_centered_line(doc, "OneRoot Rentals & Apartments", size=11, color=COLOR_MUTED, bold=True, after=2)
    add_centered_line(doc, "TENANCY AGREEMENT", size=22, color=COLOR_BLACK, bold=True, after=4)
    add_centered_line(
        doc,
        f"{PROPERTY_NAME} | Effective Upon Payment Confirmation",
        size=12,
        color=COLOR_DARK_BLUE,
        bold=True,
        after=8,
    )
    add_centered_line(
        doc,
        "Prepared from the agreed Suite 6 terms for issue after the full advance rent has been paid.",
        size=10,
        color=COLOR_MUTED,
        italic=True,
        after=14,
    )

    add_body_paragraph(
        doc,
        "This Tenancy Agreement is intended for execution once the Landlord has confirmed receipt of the full advance rent for Suite 6 - Patience. The agreement should be completed by inserting the Tenant's final details, the commencement date, and the payment confirmation details before signing.",
        after=10,
    )

    doc.add_paragraph("Agreement Snapshot", style="Heading 1")
    add_snapshot_table(doc)

    doc.add_paragraph("1. Parties", style="Heading 1")
    add_label_value_paragraph(doc, "Landlord", LANDLORD_NAME)
    for line in LANDLORD_ADDRESS_LINES:
        add_body_paragraph(doc, line, after=3)
    add_body_paragraph(doc, '(Hereinafter referred to as "the Landlord")', italic=True, after=8)
    add_label_value_paragraph(doc, "Tenant Full Legal Name", "________________________________________")
    add_label_value_paragraph(doc, "Tenant Address", "________________________________________")
    add_label_value_paragraph(doc, "Tenant Mobile", "________________________________________")
    add_label_value_paragraph(doc, "Tenant ID / Passport No.", "________________________________________")
    add_body_paragraph(doc, '(Hereinafter referred to as "the Tenant")', italic=True, after=8)

    doc.add_paragraph("2. Property and Permitted Use", style="Heading 1")
    add_body_paragraph(
        doc,
        f"The Landlord agrees to lease to the Tenant the residential apartment known as {PROPERTY_NAME}, situated at {PROPERTY_LOCATION}. The premises shall be used strictly for residential purposes only, and no commercial activity shall be carried on from the apartment without the Landlord's prior written consent.",
    )

    doc.add_paragraph("3. Lease Term and Activation", style="Heading 1")
    add_body_paragraph(
        doc,
        "The tenancy shall commence on _________________________ and shall continue for one (1) year, ending on _________________________, unless terminated earlier in accordance with this Agreement.",
    )
    add_body_paragraph(
        doc,
        f"This Agreement shall become effective only after the Landlord confirms receipt of the full one-year advance rent of {ADVANCE_RENT}. Until payment is confirmed and recorded in the Payment Confirmation section, this document remains a prepared draft for signature.",
    )

    doc.add_paragraph("4. Rent and Monthly Service Charges", style="Heading 1")
    add_body_paragraph(
        doc,
        f"The rent for the premises is {RENT_PER_MONTH} per month, payable one year in advance upon execution of this Agreement. The total advance rent due before activation is {ADVANCE_RENT}.",
    )
    add_body_paragraph(
        doc,
        f"Throughout the tenancy period, the Tenant shall also pay the following monthly utility and service charges via MTN Mobile Money to {MOMO_NUMBER}:",
        after=8,
    )
    add_fees_table(doc)
    add_body_paragraph(
        doc,
        f"The Tenant shall send proof of each monthly service-charge payment, such as a receipt or screenshot, via WhatsApp to {MOMO_NUMBER}.",
        before=8,
    )

    doc.add_paragraph("5. Safety, Alterations, and Property Care", style="Heading 1")
    add_body_paragraph(
        doc,
        "Gas cylinders shall not be stored in the bedroom or kitchen. All gas cylinders must be kept in the designated external compartment and connected to the stove from that location only.",
    )
    add_body_paragraph(
        doc,
        "The apartment includes the basic fittings presently installed, including the bath, toilet, and kitchen fixtures. The Tenant shall be responsible for any damage occurring during the tenancy beyond fair wear and tear.",
    )
    add_body_paragraph(
        doc,
        "Before moving out, the Tenant shall permit an inspection by designated service providers, including a plumber, electrician, and painter, to confirm that no outstanding repairs remain.",
    )
    add_body_paragraph(
        doc,
        "The Tenant shall not drill, paint, or make structural or cosmetic alterations to the premises without the Landlord's prior permission.",
    )

    doc.add_paragraph("6. Conduct, Guests, and House Rules", style="Heading 1")
    add_body_paragraph(
        doc,
        "No pets are allowed on the premises. The Tenant shall maintain reasonable noise levels and shall not play loud music in a manner that disturbs neighbours or other occupants.",
    )
    add_body_paragraph(
        doc,
        "Overnight guests may stay only for a limited number of days and only with the Landlord's prior approval. Smoking is not permitted inside the apartment.",
    )
    add_body_paragraph(
        doc,
        "Excessive alcohol use or any disorderly behaviour that causes nuisance, danger, or disturbance is strictly prohibited.",
    )
    add_body_paragraph(
        doc,
        "The Tenant shall keep both the apartment and the immediate surroundings in a clean and orderly condition and shall cooperate with other occupants to maintain a sanitary environment.",
    )
    add_body_paragraph(
        doc,
        "The main gate must remain locked at all times to support the security of all residents and to prevent the dogs on the premises from leaving the compound.",
    )

    doc.add_paragraph("7. Maintenance, Access, and Issue Resolution", style="Heading 1")
    add_body_paragraph(
        doc,
        "Non-urgent maintenance issues shall be reported to the Landlord via WhatsApp or text message. Urgent issues, including water leaks or electrical faults, must be reported immediately.",
    )
    add_body_paragraph(
        doc,
        "Approved service providers are as follows: Electrician - Emmanuel (contact available on the tenants' WhatsApp platform), and Plumber - Sammy (contact available on the platform). No substitute electrician shall be engaged without the Landlord's approval.",
    )
    add_body_paragraph(
        doc,
        "The Landlord may enter the premises only in an emergency, to carry out necessary or agreed repairs, inspections, or improvements, pursuant to a lawful court order, or where the Tenant is reasonably believed to have abandoned the premises.",
    )
    add_body_paragraph(
        doc,
        "The parties shall communicate respectfully and attempt to resolve issues amicably. Where a dispute remains unresolved, the parties may seek an appropriate lawful remedy, but pending resolution the Landlord's house rules shall continue to apply.",
    )

    doc.add_paragraph("8. Termination, Expiry, and Governing Law", style="Heading 1")
    add_body_paragraph(
        doc,
        "Either party may terminate this Agreement by giving not less than one (1) month's written notice, subject to any outstanding financial obligations already due under this Agreement.",
    )
    add_body_paragraph(
        doc,
        "The Tenant shall be notified when the tenancy period is approaching expiry. Any renewal or extension shall be by fresh agreement between the parties.",
    )
    add_body_paragraph(
        doc,
        "This Agreement shall be governed by and construed in accordance with the laws of the Republic of Ghana.",
    )

    doc.add_paragraph("9. Payment Confirmation", style="Heading 1")
    add_body_paragraph(
        doc,
        "Complete this section once the advance rent has been received. The agreement should be signed only after this information has been entered and confirmed by the Landlord.",
        after=8,
    )
    add_payment_confirmation_table(doc)

    doc.add_paragraph("10. Signatures", style="Heading 1")
    add_body_paragraph(
        doc,
        "IN WITNESS WHEREOF, the parties hereto have executed this Tenancy Agreement on the date written against their signatures below.",
        after=8,
    )
    add_signature_table(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
