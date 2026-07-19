from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


WORKSPACE = Path("/Users/philipboakye/Documents/OneRoot Expense Register")
SOURCE_PATH = WORKSPACE / "outputs" / "Suite_6_Patience_Tenancy_Agreement_Once_Paid.docx"
OUTPUT_PATH = WORKSPACE / "Tenancy_Agreement_Template.docx"


def clear_paragraph_runs(paragraph) -> None:
    for child in list(paragraph._element):
        if child.tag in {qn("w:r"), qn("w:hyperlink")}:
            paragraph._element.remove(child)


def copy_run_style(source_run, target_run) -> None:
    if source_run is None:
        return

    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size

    if source_run.font.color.rgb is not None:
        target_run.font.color.rgb = source_run.font.color.rgb

    if source_run._element.rPr is not None and source_run.font.name:
        target_run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), source_run.font.name)
        target_run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), source_run.font.name)


def set_paragraph_runs(paragraph, run_specs) -> None:
    source_runs = list(paragraph.runs)
    clear_paragraph_runs(paragraph)

    for style_index, text in run_specs:
        run = paragraph.add_run(text)
        source_run = source_runs[style_index] if 0 <= style_index < len(source_runs) else None
        copy_run_style(source_run, run)


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    set_paragraph_runs(paragraph, [(0, text)])


def main() -> None:
    document = Document(SOURCE_PATH)

    set_paragraph_runs(document.paragraphs[2], [(0, "[[SUITE_NAME]] | Effective Upon Payment Confirmation")])
    set_paragraph_runs(
        document.paragraphs[3],
        [(0, "Prepared from the saved apartment record for issue after the agreed advance rent has been paid.")],
    )
    set_paragraph_runs(
        document.paragraphs[4],
        [
            (
                0,
                "This Tenancy Agreement is intended for execution once the Landlord has confirmed receipt of the agreed "
                "advance rent for [[SUITE_NAME]]. The agreement should be completed by confirming the Tenant details, "
                "the commencement date, and the payment confirmation details before signing.",
            )
        ],
    )
    set_paragraph_runs(document.paragraphs[13], [(0, "Tenant Full Legal Name: "), (1, "[[TENANT_NAME]]")])
    set_paragraph_runs(document.paragraphs[14], [(0, "Tenant Address: "), (1, "[[TENANT_ADDRESS]]")])
    set_paragraph_runs(document.paragraphs[15], [(0, "Tenant Mobile: "), (1, "[[TENANT_PHONE]]")])
    set_paragraph_runs(document.paragraphs[16], [(0, "Tenant ID / Passport No.: "), (1, "[[TENANT_ID_REF]]")])
    set_paragraph_runs(
        document.paragraphs[19],
        [
            (
                0,
                "The Landlord agrees to lease to the Tenant the residential apartment known as [[SUITE_NAME]], "
                "situated at [[PROPERTY_LOCATION]]. The premises shall be used strictly for residential purposes only, "
                "and no commercial activity shall be carried on from the apartment without the Landlord's prior written consent.",
            )
        ],
    )
    set_paragraph_runs(
        document.paragraphs[21],
        [
            (
                0,
                "The tenancy shall commence on [[COMMENCEMENT_DATE]] and shall continue for [[LEASE_TERM_TEXT]], "
                "ending on [[EXPIRY_DATE]], unless terminated earlier in accordance with this Agreement.",
            )
        ],
    )
    set_paragraph_runs(
        document.paragraphs[22],
        [
            (
                0,
                "This Agreement shall become effective only after the Landlord confirms receipt of the agreed advance "
                "rent of [[ADVANCE_RENT_DUE]]. Until payment is confirmed and recorded in the Payment Confirmation "
                "section, this document remains a prepared draft for signature.",
            )
        ],
    )
    set_paragraph_runs(
        document.paragraphs[24],
        [
            (
                0,
                "The agreed rent plan is [[RENT_PLAN_SUMMARY]] The total advance rent due before activation is "
                "[[ADVANCE_RENT_DUE]].",
            )
        ],
    )
    set_paragraph_runs(
        document.paragraphs[25],
        [
            (
                0,
                "Throughout the tenancy period, the Tenant shall also pay the following monthly utility and service "
                "charges through [[PAYMENT_CHANNEL]].[[CUSTOM_BILL_ITEMS_NOTE]]",
            )
        ],
    )
    set_paragraph_runs(
        document.paragraphs[26],
        [
            (
                0,
                "The Tenant shall send proof of each monthly service-charge payment to the Landlord immediately after payment is made.",
            )
        ],
    )
    set_paragraph_runs(
        document.paragraphs[48],
        [
            (
                0,
                "Complete this section once the recorded rent payment has been received. The agreement should be signed "
                "only after this information has been entered and confirmed by the Landlord.",
            )
        ],
    )

    snapshot_table = document.tables[0]
    set_cell_text(snapshot_table.rows[0].cells[1], "[[SUITE_NAME]]")
    set_cell_text(snapshot_table.rows[1].cells[1], "[[PROPERTY_LOCATION]]")
    set_cell_text(snapshot_table.rows[2].cells[1], "[[LEASE_TERM_LABEL]]")
    set_cell_text(snapshot_table.rows[3].cells[1], "[[RENT_PLAN_SUMMARY]]")
    set_cell_text(snapshot_table.rows[4].cells[1], "[[ADVANCE_RENT_DUE]]")
    set_cell_text(snapshot_table.rows[5].cells[1], "[[MONTHLY_SERVICE_TOTAL]]")
    set_cell_text(snapshot_table.rows[6].cells[1], "[[PAYMENT_CHANNEL]]")
    set_cell_text(snapshot_table.rows[7].cells[1], "[[AGREEMENT_STATUS]]")

    fees_table = document.tables[1]
    set_cell_text(fees_table.rows[1].cells[1], "[[WATER_TOILET_BILL]]")
    set_cell_text(fees_table.rows[2].cells[1], "[[SWEEPING_BILL]]")
    set_cell_text(fees_table.rows[3].cells[1], "[[WASTE_BILL]]")
    set_cell_text(fees_table.rows[4].cells[1], "[[MONTHLY_SERVICE_TOTAL]]")

    payment_table = document.tables[2]
    set_cell_text(payment_table.rows[0].cells[1], "[[AMOUNT_RECEIVED]]")
    set_cell_text(payment_table.rows[1].cells[1], "[[PAYMENT_DATE]]")
    set_cell_text(payment_table.rows[2].cells[1], "[[PAYMENT_METHOD]]")
    set_cell_text(payment_table.rows[3].cells[1], "[[PAYMENT_REFERENCE]]")
    set_cell_text(payment_table.rows[4].cells[1], "[[RECEIVED_BY]]")
    set_cell_text(payment_table.rows[5].cells[1], "[[AGREEMENT_EFFECTIVE_FROM]]")

    signature_table = document.tables[3]
    set_cell_text(signature_table.rows[2].cells[1], "[[TENANT_NAME]]")

    document.save(OUTPUT_PATH)
    print(f"Template saved to {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
